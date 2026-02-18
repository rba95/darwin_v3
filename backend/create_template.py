"""
Script pour créer un template DAT compatible docxtpl
IMPORTANT: Pour les tableaux dynamiques, on utilise une approche différente
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """Applique une couleur de fond à une cellule"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def create_dat_template():
    doc = Document()
    
    # Style par défaut
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # ========== PAGE DE GARDE ==========
    doc.add_paragraph()
    doc.add_paragraph()
    
    title = doc.add_heading('DOSSIER D\'ARCHITECTURE TECHNIQUE', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('{{ titre_projet }}')
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0, 102, 204)
    
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Tableau d'infos générales
    info_table = doc.add_table(rows=4, cols=2)
    info_table.style = 'Table Grid'
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    info_data = [
        ('Chef de Projet', '{{ chef_projet }}'),
        ('Contact Technique', '{{ contact_tech }}'),
        ('Date', '{{ date }}'),
        ('Version', '1.0'),
    ]
    
    for i, (label, value) in enumerate(info_data):
        cell_label = info_table.rows[i].cells[0]
        cell_value = info_table.rows[i].cells[1]
        cell_label.text = label
        cell_value.text = value
        set_cell_shading(cell_label, 'E6F2FF')
        cell_label.paragraphs[0].runs[0].bold = True
    
    doc.add_page_break()
    
    # ========== 1. OBJET DU DOCUMENT ==========
    doc.add_heading('1. Objet du Document', 1)
    doc.add_paragraph('{{ objet_document }}')
    doc.add_paragraph()
    doc.add_paragraph('{{ description_doc }}')
    doc.add_paragraph()
    
    # ========== 2. ACTEURS ==========
    doc.add_heading('2. Acteurs du Projet', 1)
    
    # Pour docxtpl, la boucle doit être sur une LIGNE de tableau complète
    # On crée un tableau avec la syntaxe correcte
    actor_table = doc.add_table(rows=1, cols=4)
    actor_table.style = 'Table Grid'
    
    # En-têtes
    headers = ['Acteur', 'Rôle', 'Droits', 'Commentaires']
    for i, header in enumerate(headers):
        cell = actor_table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, '2E5090')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()
    # Utiliser la syntaxe de liste pour les acteurs
    doc.add_paragraph('{% for acteur in acteurs %}')
    doc.add_paragraph('• {{ acteur.acteur }} - {{ acteur.role }} ({{ acteur.droits }})')
    doc.add_paragraph('{% endfor %}')
    
    doc.add_paragraph()
    
    # ========== 3. ARCHITECTURE ==========
    doc.add_heading('3. Architecture Technique', 1)
    
    doc.add_heading('3.1 Description Générale', 2)
    doc.add_paragraph('{{ description_architecture }}')
    
    doc.add_heading('3.2 Authentification', 2)
    doc.add_paragraph('{{ description_authentification }}')
    
    doc.add_heading('3.3 Administration Technique', 2)
    doc.add_paragraph('{{ description_administrationtechnique }}')
    
    doc.add_heading('3.4 Administration Fonctionnelle', 2)
    doc.add_paragraph('{{ description_adminfonctionnelle }}')
    
    doc.add_heading('3.5 Communication Inter-Applicative', 2)
    doc.add_paragraph('{{ description_interapplicative }}')
    
    doc.add_page_break()
    
    # ========== 4. INFRASTRUCTURE - VMS ==========
    doc.add_heading('4. Infrastructure - Machines Virtuelles', 1)
    
    p = doc.add_paragraph()
    p.add_run('Segmentation DR : ').bold = True
    p.add_run('{{ segmentation_dr }}')
    doc.add_paragraph()
    
    # Liste des VMs
    doc.add_paragraph('{% for vm in vms %}')
    doc.add_paragraph('• {{ vm.nom }} ({{ vm.environnement }}) - {{ vm.role }}')
    doc.add_paragraph('  OS: {{ vm.os }} | CPU: {{ vm.cpu }} | RAM: {{ vm.ram }} GB')
    doc.add_paragraph('{% endfor %}')
    
    doc.add_paragraph()
    
    # ========== 5. CHOIX TECHNOLOGIQUES ==========
    doc.add_heading('5. Choix Technologiques', 1)
    
    doc.add_paragraph('{% for tech in choix_technologiques %}')
    doc.add_paragraph('• {{ tech.tiers }}: {{ tech.produit }} v{{ tech.version }}')
    doc.add_paragraph('{% endfor %}')
    
    doc.add_paragraph()
    
    # ========== 6. DNS ==========
    doc.add_heading('6. Noms DNS', 1)
    
    doc.add_paragraph('{% for dns in dns_nom %}')
    doc.add_paragraph('• {{ dns.nom_dns }} → {{ dns.machine_associe }}')
    doc.add_paragraph('{% endfor %}')
    
    doc.add_page_break()
    
    # ========== 7. CYCLE DE VIE ==========
    doc.add_heading('7. Cycle de Vie', 1)
    
    doc.add_heading('7.1 Déploiement', 2)
    doc.add_paragraph('{{ deploiement }}')
    
    doc.add_heading('7.2 Migration et Reprise', 2)
    doc.add_paragraph('{{ migration_reprise }}')
    
    doc.add_heading('7.3 Supervision', 2)
    doc.add_paragraph('{{ supervision }}')
    
    doc.add_heading('7.4 Sauvegarde et Restauration', 2)
    doc.add_paragraph('{{ sauvegarde_restauration }}')
    
    # ========== 8. DÉPENDANCES ==========
    doc.add_heading('8. Dépendances', 1)
    
    doc.add_heading('8.1 Dépendances Externes', 2)
    doc.add_paragraph('{% for dep in dependances_externes %}')
    doc.add_paragraph('• {{ dep.dependance }}: {{ dep.impact }}')
    doc.add_paragraph('{% endfor %}')
    
    doc.add_heading('8.2 Dépendances vers Services Externes', 2)
    doc.add_paragraph('{% for dep in dependance_app_externes %}')
    doc.add_paragraph('• {{ dep.name_application }}: {{ dep.description_impact }}')
    doc.add_paragraph('{% endfor %}')
    
    doc.add_paragraph()
    
    # ========== 9. CONTRAINTES & SLA ==========
    doc.add_heading('9. Contraintes et Niveau de Service', 1)
    
    doc.add_heading('9.1 Contraintes', 2)
    doc.add_paragraph('{{ contraintes }}')
    
    doc.add_heading('9.2 Niveau de Services (SLA)', 2)
    doc.add_paragraph('{{ niveau_services }}')
    
    # ========== SAUVEGARDE ==========
    output_path = 'app/templates/dat_template.docx'
    doc.save(output_path)
    print(f"✅ Template DAT créé : {output_path}")
    return output_path

if __name__ == "__main__":
    create_dat_template()
